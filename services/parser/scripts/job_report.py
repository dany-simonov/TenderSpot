import os
import re
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Iterable, List, Optional

import requests
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

DOWNLOADS_DIR = Path(r"C:\Users\redmi\Downloads")
DATE_FORMAT = "%d.%m.%Y"
FILE_DATE_FORMAT = "%Y-%m-%d"

ENV_PATHS = [
    Path.cwd() / ".env",
    Path(__file__).resolve().parents[3] / ".env",
]


@dataclass
class ReportItem:
    title: str
    location: str
    budget: str
    deadline: str
    url: str


def _normalize(text: Optional[str]) -> str:
    if not text:
        return "—"
    return re.sub(r"\s+", " ", text).strip()


def _load_env_file(path: Path) -> None:
    if not path.exists():
        return
    for line in path.read_text(encoding="utf-8").splitlines():
        stripped = line.strip()
        if not stripped or stripped.startswith("#") or "=" not in stripped:
            continue
        key, value = stripped.split("=", 1)
        key = key.strip()
        value = value.strip().strip('"').strip("'")
        os.environ.setdefault(key, value)


def _ensure_env() -> None:
    for path in ENV_PATHS:
        _load_env_file(path)


def _format_budget(value: Optional[object]) -> str:
    if value is None:
        return "—"
    try:
        amount = float(value)
    except (TypeError, ValueError):
        return _normalize(str(value))
    if amount.is_integer():
        return f"{int(amount):,} RUB".replace(",", " ")
    return f"{amount:,.2f} RUB".replace(",", " ")


def fetch_items() -> List[ReportItem]:
    _ensure_env()

    endpoint = os.getenv("APPWRITE_ENDPOINT") or os.getenv("VITE_APPWRITE_ENDPOINT")
    project_id = os.getenv("APPWRITE_PROJECT_ID") or os.getenv("VITE_APPWRITE_PROJECT_ID")
    api_key = os.getenv("PARSER_APPWRITE_API_KEY")
    database_id = os.getenv("APPWRITE_DATABASE_ID") or os.getenv("VITE_APPWRITE_DATABASE_ID")
    collection_id = (
        os.getenv("APPWRITE_TENDERS_COLLECTION_ID")
        or os.getenv("VITE_APPWRITE_TENDERS_COLLECTION_ID")
        or os.getenv("VITE_APPWRITE_COLLECTION_ID")
    )

    if not endpoint or not project_id or not api_key or not database_id or not collection_id:
        raise ValueError(
            "Appwrite env is incomplete. Required: APPWRITE_ENDPOINT, APPWRITE_PROJECT_ID, "
            "PARSER_APPWRITE_API_KEY, APPWRITE_DATABASE_ID, APPWRITE_TENDERS_COLLECTION_ID."
        )

    limit = int(os.getenv("REPORT_LIMIT", "15"))
    base_url = endpoint.rstrip("/")
    url = f"{base_url}/databases/{database_id}/collections/{collection_id}/documents"
    params = [
        ("queries[]", f"limit({limit})"),
        ("queries[]", "orderDesc($createdAt)"),
    ]
    headers = {
        "X-Appwrite-Project": project_id,
        "X-Appwrite-Key": api_key,
    }

    response = requests.get(url, params=params, headers=headers, timeout=30)
    response.raise_for_status()
    payload = response.json()
    documents = payload.get("documents", [])

    items: List[ReportItem] = []
    for doc in documents:
        title = _normalize(doc.get("title"))
        location = _normalize(doc.get("regionCode"))
        budget = _format_budget(doc.get("price"))
        deadline = _normalize(doc.get("deadline"))
        url_value = _normalize(doc.get("url"))

        if title == "—" and url_value == "—":
            continue

        items.append(
            ReportItem(
                title=title,
                location=location,
                budget=budget,
                deadline=deadline,
                url=url_value,
            )
        )

    return items


def _add_hyperlink(paragraph, text: str, url: str) -> None:
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    new_run.append(r_pr)

    text_node = OxmlElement("w:t")
    text_node.text = text
    new_run.append(text_node)
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)


def _ensure_unique_path(path: Path) -> Path:
    if not path.exists():
        return path

    counter = 2
    while True:
        candidate = path.with_stem(f"{path.stem}_v{counter}")
        if not candidate.exists():
            return candidate
        counter += 1


def build_report(items: Iterable[ReportItem]) -> Path:
    now = datetime.now()
    document = Document()

    style = document.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    document.add_heading(f"Отчет от {now.strftime(DATE_FORMAT)}", level=1)

    for idx, item in enumerate(items, start=1):
        title_paragraph = document.add_paragraph()
        title_run = title_paragraph.add_run(f"{idx}. {item.title}")
        title_run.bold = True

        document.add_paragraph(f"📍 Место: {item.location}")
        document.add_paragraph(f"💰 Бюджет: {item.budget}")
        document.add_paragraph(f"⏰ Дедлайн: {item.deadline}")

        link_paragraph = document.add_paragraph("🔗 Ссылка: ")
        if item.url and item.url != "—":
            _add_hyperlink(link_paragraph, item.url, item.url)
        else:
            link_paragraph.add_run("—")

        document.add_paragraph("")

    filename = f"report_{now.strftime(FILE_DATE_FORMAT)}.docx"
    target_path = _ensure_unique_path(DOWNLOADS_DIR / filename)
    document.save(target_path)
    return target_path


def main() -> None:
    items = fetch_items()
    if not items:
        raise RuntimeError("No items parsed. Check selectors and source URL.")

    output_path = build_report(items)
    print(f"Report saved to: {output_path}")


if __name__ == "__main__":
    main()
